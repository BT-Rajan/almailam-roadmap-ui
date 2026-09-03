"""Admin-uploaded .docx templates for Quotation/Contract documents.

An Administrator uploads a .docx per document_type (Quotation, Contract)
under Administration > Documents and marks one as the default. When a
project's Quotation/Contract tab asks for the actual document, the
default template's placeholders ({{ field }}, and {%tr for ... %} row
loops inside a table -- both standard docxtpl/Jinja2 syntax) are merged
with that record's live data via render_quotation_document/
render_contract_document below, and the merged .docx is the download.

render_quotation_pdf/render_contract_pdf below convert that same merged
.docx to PDF (via _docx_to_pdf) so "Print" and "Email" can use the
identical, admin-configured template instead of the separate hardcoded
on-screen preview (QuotationPreview.vue/ContractPreview.vue) those
actions used to print.
"""

import copy
import html
import io
import re
from datetime import datetime, timezone
from decimal import Decimal

import mammoth
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate
from sqlalchemy.orm import Session
from weasyprint import HTML

from app.core.exceptions import NotFoundError, ValidationAppError
from app.core.file_storage import resolve_path, save_upload
from app.core.number_to_words import amount_to_words
from app.models.client import Client
from app.models.contract import Contract
from app.models.document_template import DOCUMENT_TEMPLATE_TYPES, DocumentTemplate
from app.models.project import Project
from app.models.quotation import Quotation
from app.models.user import User
from app.services import audit_service
from app.services.pdf_render import FONT_PATH

ENTITY_TYPE = "DOCUMENT_TEMPLATE"
STORAGE_SUBDIRECTORY = "document-templates"

_TAG_RE = re.compile(r"<[^>]+>")
_BLOCK_BREAK_RE = re.compile(r"</(p|div|li|br)\s*/?>", re.IGNORECASE)


def _plain_text(value: str | None) -> str:
    """Strips the small rich-text HTML allowlist (see
    core/html_sanitizer.py) down to plain text for a Word merge field --
    a .docx template has no way to render arbitrary saved HTML, only
    plain text/Jinja2 placeholders. Block-level tags become newlines so
    paragraphs/list items don't run together; everything else is
    stripped outright (inline images included -- they don't survive the
    merge, only their surrounding text does)."""
    if not value:
        return ""
    with_breaks = _BLOCK_BREAK_RE.sub("\n", value)
    stripped = _TAG_RE.sub("", with_breaks)
    return html.unescape(stripped).strip()


def _user_name(db: Session, user_id: int) -> str:
    user = db.query(User).filter(User.id == user_id).first()
    return user.full_name if user else "Unknown"


def _check_document_type(document_type: str) -> None:
    if document_type not in DOCUMENT_TEMPLATE_TYPES:
        raise ValidationAppError(f"documentType must be one of {DOCUMENT_TEMPLATE_TYPES}")


def list_templates(db: Session, document_type: str | None = None) -> list[DocumentTemplate]:
    query = db.query(DocumentTemplate).filter(DocumentTemplate.deleted_at.is_(None))
    if document_type is not None:
        _check_document_type(document_type)
        query = query.filter(DocumentTemplate.document_type == document_type)
    return query.order_by(DocumentTemplate.id.desc()).all()


def get_template(db: Session, template_id: int) -> DocumentTemplate:
    template = (
        db.query(DocumentTemplate)
        .filter(DocumentTemplate.id == template_id, DocumentTemplate.deleted_at.is_(None))
        .first()
    )
    if template is None:
        raise NotFoundError("Document template")
    return template


def get_default(db: Session, document_type: str) -> DocumentTemplate | None:
    _check_document_type(document_type)
    return (
        db.query(DocumentTemplate)
        .filter(
            DocumentTemplate.document_type == document_type,
            DocumentTemplate.is_default.is_(True),
            DocumentTemplate.deleted_at.is_(None),
        )
        .first()
    )


def upload_template(db: Session, document_type: str, file, actor_id: int) -> DocumentTemplate:
    _check_document_type(document_type)
    if not (file.filename or "").lower().endswith(".docx"):
        raise ValidationAppError("Only Word (.docx) files are accepted for a document template.")

    storage_key, original_filename, size_bytes = save_upload(file, STORAGE_SUBDIRECTORY)

    template = DocumentTemplate(
        document_type=document_type,
        storage_key=storage_key,
        original_filename=original_filename,
        file_size_bytes=size_bytes,
        uploaded_by=actor_id,
        # The first template uploaded for a type becomes its default
        # automatically -- otherwise "Download Document" would 404 with
        # no default configured until an admin remembers to set one.
        is_default=get_default(db, document_type) is None,
    )
    db.add(template)
    db.flush()
    audit_service.log_event(
        db, ENTITY_TYPE, template.id, f"{document_type} template uploaded", actor_id, new_value=original_filename
    )
    db.commit()
    db.refresh(template)
    return template


def set_default(db: Session, template_id: int, actor_id: int) -> DocumentTemplate:
    template = get_template(db, template_id)
    if template.is_default:
        return template
    db.query(DocumentTemplate).filter(
        DocumentTemplate.document_type == template.document_type,
        DocumentTemplate.is_default.is_(True),
    ).update({"is_default": False})
    template.is_default = True
    audit_service.log_event(
        db, ENTITY_TYPE, template.id, f"{template.document_type} default template changed", actor_id,
        new_value=template.original_filename,
    )
    db.commit()
    db.refresh(template)
    return template


def delete_template(db: Session, template_id: int, actor_id: int) -> None:
    template = get_template(db, template_id)
    if template.is_default:
        raise ValidationAppError(
            "This is the default template for its document type. Set another template as default first."
        )
    audit_service.log_event(
        db, ENTITY_TYPE, template.id, f"{template.document_type} template deleted", actor_id,
        previous_value=template.original_filename,
    )
    template.deleted_at = datetime.now(timezone.utc)
    db.commit()


def _render_docx(storage_key: str, context: dict) -> bytes:
    path = resolve_path(storage_key)
    try:
        doc = DocxTemplate(str(path))
        doc.render(context)
    except Exception as exc:
        raise ValidationAppError(
            f"The uploaded template couldn't be merged -- check its placeholder syntax. ({exc})"
        ) from exc
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Converts a merged .docx to a real, selectable-text PDF for
    Print/Email, reusing the exact stack pdf_render.py already relies on
    for Government Forms (WeasyPrint + the bundled Noto Naskh Arabic
    font) rather than adding a new system dependency: mammoth (a pure
    pip package, no system libs) turns the .docx body into semantic
    HTML, which WeasyPrint then paginates into a PDF. This won't
    reproduce the admin's exact Word layout pixel-for-pixel (mammoth
    keeps paragraphs/headings/bold/tables/lists, not precise
    spacing/borders), but it is the same merged content, real
    RTL-capable Arabic text throughout -- unlike a LibreOffice-based
    docx->PDF conversion, which was tried here and does not run
    reliably headless in this app's container."""
    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    body_html = result.value
    html_doc = f"""<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
<meta charset="utf-8">
<style>
  @font-face {{
    font-family: 'NotoNaskhArabic';
    src: url('file://{FONT_PATH}');
  }}
  @page {{ size: A4; margin: 2.5cm 2cm; }}
  body {{
    font-family: 'NotoNaskhArabic', sans-serif;
    direction: rtl;
    text-align: right;
    font-size: 12pt;
    line-height: 1.8;
    color: #1a1a2e;
  }}
  table {{ width: 100%; border-collapse: collapse; margin: 0 0 1em; }}
  td, th {{ border: 1px solid #999; padding: 0.4em 0.6em; }}
  p {{ margin: 0 0 0.8em; }}
</style>
</head>
<body>{body_html}</body>
</html>"""
    return HTML(string=html_doc).write_pdf()


def render_quotation_document(db: Session, quotation: Quotation) -> tuple[bytes, str]:
    from app.services import quotation_service

    template = get_default(db, "Quotation")
    if template is None:
        raise ValidationAppError(
            "No default Quotation template is configured. Upload one in Administration > Documents."
        )

    project = db.query(Project).filter(Project.id == quotation.project_id).first()
    client = (
        db.query(Client).filter(Client.id == project.client_id).first()
        if project is not None
        else None
    )
    line_items = quotation_service.get_line_items(db, quotation.id)
    subtotal = sum((Decimal(str(i.quantity)) * Decimal(str(i.unit_price)) for i in line_items), Decimal("0"))

    context = {
        "quotation_no": quotation.quotation_no,
        "revision": quotation.revision,
        "issue_date": quotation.issue_date.strftime("%d %B %Y"),
        "validity": quotation.validity.strftime("%d %B %Y"),
        "status": quotation.status,
        "currency": quotation.currency,
        "prepared_by": _user_name(db, quotation.prepared_by),
        "client_name": client.company_name if client else "",
        "project_name": project.project_name if project else "",
        "project_no": project.project_no if project else "",
        "project_address": (project.site_address or "") if project else "",
        "amount_in_words": amount_to_words(Decimal(str(quotation.amount)), quotation.currency),
        "line_items": [
            {
                "description": item.description,
                "quantity": f"{float(item.quantity):g}",
                "unit_price": f"{float(item.unit_price):.2f}",
                "amount": f"{float(item.quantity) * float(item.unit_price):.2f}",
            }
            for item in line_items
        ],
        "subtotal": f"{subtotal:.2f}",
        "discount_amount": f"{float(quotation.discount_amount):.2f}",
        "amount": f"{float(quotation.amount):.2f}",
        "notes": _plain_text(quotation.notes),
        "terms_and_conditions": [_plain_text(term) for term in quotation.terms_and_conditions],
        "scope_phases": [_plain_text(phase) for phase in quotation.scope_phases],
        "payment_terms": [_plain_text(term) for term in quotation.payment_terms],
    }
    filename = f"{quotation.quotation_no}.docx"
    return _render_docx(template.storage_key, context), filename


def render_quotation_pdf(db: Session, quotation: Quotation) -> tuple[bytes, str]:
    """Same merged document as render_quotation_document, converted to
    PDF -- what Print and Email actually use, so both show the admin's
    real uploaded template rather than a separate hardcoded preview."""
    content, filename = render_quotation_document(db, quotation)
    return _docx_to_pdf(content), filename.removesuffix(".docx") + ".pdf"


def render_contract_document(db: Session, contract: Contract) -> tuple[bytes, str]:
    from app.services import contract_service

    template = get_default(db, "Contract")
    if template is None:
        raise ValidationAppError(
            "No default Contract template is configured. Upload one in Administration > Documents."
        )

    project = db.query(Project).filter(Project.id == contract.project_id).first()
    client = (
        db.query(Client).filter(Client.id == project.client_id).first()
        if project is not None
        else None
    )
    clauses = contract_service.get_clauses(db, contract.id)

    context = {
        "contract_no": contract.contract_no,
        "revision": contract.revision,
        "currency": contract.currency,
        "contract_value": f"{float(contract.contract_value):.2f}",
        "issue_date": contract.issue_date.strftime("%d %B %Y"),
        "signed_date": contract.signed_date.strftime("%d %B %Y") if contract.signed_date else "",
        "expiry_date": contract.expiry_date.strftime("%d %B %Y"),
        "status": contract.status,
        "prepared_by": _user_name(db, contract.prepared_by),
        "client_representative": contract.client_representative,
        "client_name": client.company_name if client else "",
        "project_name": project.project_name if project else "",
        "project_no": project.project_no if project else "",
        "project_address": (project.site_address or "") if project else "",
        "amount_in_words": amount_to_words(Decimal(str(contract.contract_value)), contract.currency),
        "scope_summary": _plain_text(contract.scope_summary),
        "clauses": [{"title": c.title, "content": _plain_text(c.content)} for c in clauses],
    }
    filename = f"{contract.contract_no}.docx"
    return _render_docx(template.storage_key, context), filename


def render_contract_pdf(db: Session, contract: Contract) -> tuple[bytes, str]:
    """PDF counterpart of render_contract_document -- see
    render_quotation_pdf's docstring."""
    content, filename = render_contract_document(db, contract)
    return _docx_to_pdf(content), filename.removesuffix(".docx") + ".pdf"


# --- Visual field mapping -- lets an admin click a spot in an uploaded
# template and place a merge field there instead of hand-typing {{ field
# }}/{%tr %} syntax into Word. The catalog below is the single source of
# truth for what a template can reference; it must stay in lockstep with
# the context dicts render_quotation_document/render_contract_document
# actually build above, since _sample_context's save-time validation
# render is only as meaningful as that agreement.
MERGE_FIELD_CATALOG: dict[str, list[dict]] = {
    "Quotation": [
        {"key": "quotation_no", "label": "Quotation No.", "kind": "text"},
        {"key": "revision", "label": "Revision", "kind": "text"},
        {"key": "issue_date", "label": "Issue Date", "kind": "text"},
        {"key": "validity", "label": "Valid Until", "kind": "text"},
        {"key": "status", "label": "Status", "kind": "text"},
        {"key": "currency", "label": "Currency", "kind": "text"},
        {"key": "prepared_by", "label": "Prepared By", "kind": "text"},
        {"key": "client_name", "label": "Client Name", "kind": "text"},
        {"key": "project_name", "label": "Project Name", "kind": "text"},
        {"key": "project_no", "label": "Project No.", "kind": "text"},
        {"key": "project_address", "label": "Project/Site Address", "kind": "text"},
        {"key": "subtotal", "label": "Subtotal", "kind": "text"},
        {"key": "discount_amount", "label": "Discount Amount", "kind": "text"},
        {"key": "amount", "label": "Total Amount", "kind": "text"},
        {"key": "amount_in_words", "label": "Total Amount (in Words)", "kind": "text"},
        {"key": "notes", "label": "Notes", "kind": "text"},
        {
            "key": "line_items",
            "label": "Line Items",
            "kind": "repeating_table",
            "loopVar": "item",
            "columns": [
                {"key": "description", "label": "Description"},
                {"key": "quantity", "label": "Quantity"},
                {"key": "unit_price", "label": "Unit Price"},
                {"key": "amount", "label": "Amount"},
            ],
        },
        {"key": "terms_and_conditions", "label": "Terms & Conditions", "kind": "repeating_list", "loopVar": "term"},
        {"key": "scope_phases", "label": "Scope Phases", "kind": "repeating_list", "loopVar": "phase"},
        {"key": "payment_terms", "label": "Payment Terms", "kind": "repeating_list", "loopVar": "term"},
    ],
    "Contract": [
        {"key": "contract_no", "label": "Contract No.", "kind": "text"},
        {"key": "revision", "label": "Revision", "kind": "text"},
        {"key": "currency", "label": "Currency", "kind": "text"},
        {"key": "contract_value", "label": "Contract Value", "kind": "text"},
        {"key": "issue_date", "label": "Issue Date", "kind": "text"},
        {"key": "signed_date", "label": "Signed Date", "kind": "text"},
        {"key": "expiry_date", "label": "Expiry Date", "kind": "text"},
        {"key": "status", "label": "Status", "kind": "text"},
        {"key": "prepared_by", "label": "Prepared By", "kind": "text"},
        {"key": "client_representative", "label": "Client Representative", "kind": "text"},
        {"key": "client_name", "label": "Client Name", "kind": "text"},
        {"key": "project_name", "label": "Project Name", "kind": "text"},
        {"key": "project_no", "label": "Project No.", "kind": "text"},
        {"key": "project_address", "label": "Project/Site Address", "kind": "text"},
        {"key": "amount_in_words", "label": "Contract Value (in Words)", "kind": "text"},
        {"key": "scope_summary", "label": "Scope Summary", "kind": "text"},
        {
            "key": "clauses",
            "label": "Clauses",
            "kind": "repeating_table",
            "loopVar": "clause",
            "columns": [
                {"key": "title", "label": "Title"},
                {"key": "content", "label": "Content"},
            ],
        },
    ],
}

_FIELD_BY_KEY: dict[str, dict[str, dict]] = {
    doc_type: {field["key"]: field for field in fields} for doc_type, fields in MERGE_FIELD_CATALOG.items()
}


def get_merge_fields(document_type: str) -> list[dict]:
    _check_document_type(document_type)
    return MERGE_FIELD_CATALOG[document_type]


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Rewrites a paragraph down to a single run containing `text`,
    preserving the first existing run's character formatting (bold,
    font, etc.) if there was one. A deliberate simplification -- a
    paragraph that mixed multiple run styles mid-sentence collapses to
    one style -- far more robust than trying to splice text in while
    preserving every original run boundary, and irrelevant for the
    short, mostly-plain lines a merge field actually lives in."""
    runs = list(paragraph.runs)
    preserved_rpr = None
    if runs:
        existing_rpr = runs[0]._r.find(qn("w:rPr"))
        if existing_rpr is not None:
            preserved_rpr = copy.deepcopy(existing_rpr)
    for run in runs:
        run._r.getparent().remove(run._r)
    new_run = OxmlElement("w:r")
    if preserved_rpr is not None:
        new_run.append(preserved_rpr)
    new_text_el = OxmlElement("w:t")
    new_text_el.set(qn("xml:space"), "preserve")
    new_text_el.text = text
    new_run.append(new_text_el)
    paragraph._p.append(new_run)


def _marker_paragraph_element(tag_text: str):
    """A bare, unstyled paragraph holding one docxtpl {%p %}/{%tr %} tag
    -- used as a disposable before/after marker, never seen by anyone:
    docxtpl's preprocessing deletes the entire <w:p>/<w:tr> a tag like
    this is found in and replaces it with just the bare Jinja tag (see
    the module-level docxtpl docs), so this element's own formatting
    never survives to be seen."""
    p = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = tag_text
    run.append(text_el)
    p.append(run)
    return p


def _clone_marker_row(row, marker_text: str):
    """A row-shaped clone of `row` (same cell count/widths/borders) with
    every cell's text cleared and one docxtpl {%tr %} marker tag written
    into its first cell. Needs a whole extra row rather than just text
    inside the marked row itself: docxtpl deletes the ENTIRE <w:tr> a
    {%tr %} tag is found in, replacing it with the bare Jinja tag -- the
    marked row itself is left as an ordinary row and is what actually
    repeats, so the for/endfor tags need their own throwaway rows
    immediately before/after it, not to live inside it."""
    cloned_element = copy.deepcopy(row._tr)
    cloned_row = type(row)(cloned_element, row._parent)
    for cell in cloned_row.cells:
        for paragraph in cell.paragraphs:
            _set_paragraph_text(paragraph, "")
    if cloned_row.cells and cloned_row.cells[0].paragraphs:
        _set_paragraph_text(cloned_row.cells[0].paragraphs[0], marker_text)
    return cloned_element


def _body_blocks(document) -> list[tuple[int, str, object]]:
    """(blockIndex, kind, wrapper) for every top-level paragraph/table in
    the document body, in document order -- extract_layout and
    apply_mapping both walk this exact same way, so a blockIndex always
    means the same block between the two calls (nothing changes the
    docx's own body structure in between; only text within existing
    blocks/cells does)."""
    blocks: list[tuple[int, str, object]] = []
    index = 0
    for child in document.element.body:
        if child.tag == qn("w:p"):
            blocks.append((index, "paragraph", Paragraph(child, document)))
            index += 1
        elif child.tag == qn("w:tbl"):
            blocks.append((index, "table", Table(child, document)))
            index += 1
    return blocks


def extract_layout(storage_key: str) -> dict:
    """The uploaded template's paragraphs/tables, in document order, as
    plain editable text -- what the admin's field-mapping screen shows
    and edits. Only a table cell's first paragraph is exposed (a cell
    with more than one paragraph is the rare case, and editing just the
    first is a reasonable simplification for a merge-field target)."""
    document = Document(str(resolve_path(storage_key)))
    blocks: list[dict] = []
    for block_index, kind, wrapper in _body_blocks(document):
        if kind == "paragraph":
            blocks.append({"kind": "paragraph", "blockIndex": block_index, "text": wrapper.text, "repeatingField": None})
        else:
            rows = []
            for row_index, row in enumerate(wrapper.rows):
                cells = [
                    {"cellIndex": cell_index, "text": cell.paragraphs[0].text if cell.paragraphs else ""}
                    for cell_index, cell in enumerate(row.cells)
                ]
                rows.append({"rowIndex": row_index, "cells": cells, "repeatingField": None})
            blocks.append({"kind": "table", "blockIndex": block_index, "rows": rows})
    return {"blocks": blocks}


def apply_mapping(storage_key: str, document_type: str, blocks: list[dict]) -> bytes:
    """Writes the admin's edited text back into the template (one plain
    run per paragraph/cell -- see _set_paragraph_text), then, for at
    most one paragraph and one table row across the whole document,
    wraps it in docxtpl's repeating-block marker tags if it was flagged
    `repeatingField` -- see _clone_marker_row/_marker_paragraph_element
    for why that needs extra rows/paragraphs rather than editing in
    place. A second block flagged for the same or another repeating
    field is simply ignored (last one processed in document order wins
    for a given kind) rather than erroring -- the frontend's palette
    only ever lets one location hold a given repeating field at a time,
    so this only matters for a payload assembled by hand."""
    _check_document_type(document_type)
    document = Document(str(resolve_path(storage_key)))
    by_index = {index: (kind, wrapper) for index, kind, wrapper in _body_blocks(document)}

    repeating_paragraph: Paragraph | None = None
    repeating_paragraph_field: dict | None = None
    repeating_row = None
    repeating_row_field: dict | None = None

    for block in blocks:
        located = by_index.get(block.get("blockIndex"))
        if located is None:
            continue
        kind, wrapper = located
        if kind != block.get("kind"):
            continue

        if kind == "paragraph":
            _set_paragraph_text(wrapper, block.get("text") or "")
            field_key = block.get("repeatingField")
            if field_key and field_key in _FIELD_BY_KEY[document_type]:
                repeating_paragraph = wrapper
                repeating_paragraph_field = _FIELD_BY_KEY[document_type][field_key]
        else:
            for row_payload in block.get("rows") or []:
                row_index = row_payload.get("rowIndex")
                if row_index is None or row_index >= len(wrapper.rows):
                    continue
                row = wrapper.rows[row_index]
                for cell_payload in row_payload.get("cells") or []:
                    cell_index = cell_payload.get("cellIndex")
                    if cell_index is None or cell_index >= len(row.cells):
                        continue
                    cell = row.cells[cell_index]
                    if cell.paragraphs:
                        _set_paragraph_text(cell.paragraphs[0], cell_payload.get("text") or "")
                field_key = row_payload.get("repeatingField")
                if field_key and field_key in _FIELD_BY_KEY[document_type]:
                    repeating_row = row
                    repeating_row_field = _FIELD_BY_KEY[document_type][field_key]

    if repeating_paragraph is not None and repeating_paragraph_field is not None:
        loop_var = repeating_paragraph_field["loopVar"]
        list_key = repeating_paragraph_field["key"]
        repeating_paragraph._p.addprevious(_marker_paragraph_element(f"{{%p for {loop_var} in {list_key} %}}"))
        repeating_paragraph._p.addnext(_marker_paragraph_element("{%p endfor %}"))

    if repeating_row is not None and repeating_row_field is not None:
        loop_var = repeating_row_field["loopVar"]
        list_key = repeating_row_field["key"]
        repeating_row._tr.addprevious(_clone_marker_row(repeating_row, f"{{%tr for {loop_var} in {list_key} %}}"))
        repeating_row._tr.addnext(_clone_marker_row(repeating_row, "{%tr endfor %}"))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _sample_context(document_type: str) -> dict:
    """A minimal but structurally-complete context (one sample row/term)
    for every field in the catalog -- used only to validate a mapping
    renders cleanly (see save_mapping) before it's allowed to overwrite
    the template's actual file. Not used for a real document -- that's
    render_quotation_document/render_contract_document, with the
    project's real data."""
    context: dict = {}
    for field in MERGE_FIELD_CATALOG[document_type]:
        if field["kind"] == "text":
            context[field["key"]] = "Sample"
        elif field["kind"] == "repeating_table":
            context[field["key"]] = [{column["key"]: "Sample" for column in field["columns"]}]
        elif field["kind"] == "repeating_list":
            context[field["key"]] = ["Sample"]
    return context


def save_mapping(db: Session, template_id: int, blocks: list[dict], actor_id: int) -> DocumentTemplate:
    template = get_template(db, template_id)
    new_bytes = apply_mapping(template.storage_key, template.document_type, blocks)

    # Fail the save, not the next real download -- catches a stray
    # {{ / {% left over from a bad edit, or two locations flagged for
    # the same repeating field, immediately, with the template's
    # previous, still-working file left untouched on disk either way.
    try:
        DocxTemplate(io.BytesIO(new_bytes)).render(_sample_context(template.document_type))
    except Exception as exc:
        raise ValidationAppError(
            f"This field mapping couldn't be validated -- check for overlapping or malformed placeholders. ({exc})"
        ) from exc

    resolve_path(template.storage_key).write_bytes(new_bytes)
    template.file_size_bytes = len(new_bytes)
    audit_service.log_event(db, ENTITY_TYPE, template.id, f"{template.document_type} template fields mapped", actor_id)
    db.commit()
    db.refresh(template)
    return template
